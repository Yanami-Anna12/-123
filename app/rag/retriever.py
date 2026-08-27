"""RAG（检索增强生成）核心：把文档变成可检索的知识库。

完整流程（一条流水线）：
    文档文件 → 解析成文本 → 切分成片段 → 向量化 → 存进 Milvus
    提问     → 向量化 → 稠密检索 → 返回 Top-K

外部包：
- 文档解析：pypdf / python-docx
- 文本切割：langchain 的 RecursiveCharacterTextSplitter
- 向量化：sentence-transformers（本地 bge-base-zh-v1.5，768 维稠密向量）
- 重排：sentence-transformers 的 CrossEncoder（可选）
- 存储检索：pymilvus（Milvus）
"""
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pymilvus import DataType, MilvusClient

from app.config.settings import settings as config

# 模型懒加载缓存：第一次真正用的时候才加载，之后复用（CPU 上加载一次要 20 多秒）
_embedding_model = None
_reranker_model = None


# ============ 1. 文档解析：文件 → 纯文本 ============

def _load_one(path):
    """读取单个文档，返回 langchain 的 Document（正文 + 来源信息）。"""
    p = Path(path)
    ext = p.suffix.lower()
    if ext in {".md", ".markdown", ".txt"}:
        # 纯文本，直接读出来
        text = p.read_text(encoding="utf-8", errors="ignore")
    elif ext == ".pdf":
        # PDF 用 pypdf 抽取每一页的文字
        from pypdf import PdfReader

        text = "\n\n".join(page.extract_text() or "" for page in PdfReader(str(p)).pages)
    elif ext == ".docx":
        # Word 用 python-docx 抽取段落文字
        import docx

        text = "\n".join(para.text for para in docx.Document(str(p)).paragraphs)
    else:
        raise ValueError(f"暂不支持的文件格式: {ext}")
    return Document(page_content=text, metadata={"source": p.name})


def load_documents(paths):
    """批量读取多个文档。"""
    return [_load_one(p) for p in paths]


# ============ 2. 文本切割：文本 → 片段（chunk） ============

def split_documents(documents):
    """用 langchain 的切割器切分文档（中文按句末标点切，段间留重叠）。"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        # 按这些分隔符从大到小尝试切：段落 → 换行 → 句末标点 → 逗号 → 空格 → 字符
        separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
    )
    return splitter.split_documents(documents)


# ============ 3. 向量化：文本 → 768 维稠密向量 ============

def get_embedding_model():
    """懒加载 bge-base-zh-v1.5（CPU）。第一次调用才加载，之后复用。"""
    global _embedding_model
    if _embedding_model is None:
        from sentence_transformers import SentenceTransformer

        _embedding_model = SentenceTransformer(config.EMBEDDING_MODEL_PATH, device="cpu")
    return _embedding_model


def _embed_texts(texts):
    """把一批文本变成 768 维稠密向量（归一化）。"""
    model = get_embedding_model()
    vecs = model.encode(
        texts,
        batch_size=8,
        normalize_embeddings=True,
        convert_to_numpy=True,
    )
    return [{"dense": [float(v) for v in vec]} for vec in vecs]


# ============ 4. 重排：Reranker 对候选片段做二次精排（可选） ============

def get_reranker_model():
    """懒加载 BGE-Reranker 模型（CPU）。"""
    global _reranker_model
    if _reranker_model is None:
        from sentence_transformers import CrossEncoder

        _reranker_model = CrossEncoder(config.RERANKER_MODEL_PATH, device="cpu")
    return _reranker_model


def _rerank(query, candidates, top_k):
    """对候选片段重新打分排序；reranker 不可用时返回 None（跳过精排）。"""
    try:
        model = get_reranker_model()
        ranked = model.rank(
            query, candidates,
            top_k=min(top_k, len(candidates)),
            convert_to_numpy=True,
        )
        return [(int(item["corpus_id"]), float(item["score"])) for item in ranked]
    except Exception:
        return None


# ============ 5. Milvus：建表 / 写入 / 稠密检索 ============

def _client():
    """连接 Milvus。"""
    return MilvusClient(uri=config.MILVUS_URI)


def _init_collection():
    """建 Milvus 集合（如果不存在）。字段：id / text / source / dense_vector。"""
    c = _client()
    if c.has_collection(config.MILVUS_COLLECTION):
        return c
    schema = c.create_schema(auto_id=True, enable_dynamic_field=False)
    schema.add_field("id", DataType.INT64, is_primary=True)
    schema.add_field("text", DataType.VARCHAR, max_length=65535)
    schema.add_field("source", DataType.VARCHAR, max_length=512)
    schema.add_field("dense_vector", DataType.FLOAT_VECTOR, dim=768)  # bge-base-zh 768 维
    index = c.prepare_index_params()
    index.add_index(field_name="dense_vector", index_type="AUTOINDEX", metric_type="COSINE")
    c.create_collection(config.MILVUS_COLLECTION, schema=schema, index_params=index)
    c.load_collection(config.MILVUS_COLLECTION)
    return c


def ingest(paths):
    """文档入库：解析 → 切割 → 向量化 → 写入 Milvus，返回写入的片段数。"""
    chunks = split_documents(load_documents(paths))
    if not chunks:
        return 0
    embeds = _embed_texts([c.page_content for c in chunks])
    c = _init_collection()
    rows = [
        {
            "text": ch.page_content,
            "source": ch.metadata.get("source", ""),
            "dense_vector": e["dense"],
        }
        for ch, e in zip(chunks, embeds)
    ]
    c.insert(config.MILVUS_COLLECTION, data=rows)
    c.flush(config.MILVUS_COLLECTION)
    return len(rows)


def search(query, top_k=None):
    """检索两步：① 向量化问题 ② 稠密检索召回 ③（可选）重排取 Top-K。"""
    top_k = top_k or config.TOP_K

    # ① 把问题变成向量
    q = _embed_texts([query])[0]

    # ② 稠密检索
    c = _client()
    results = c.search(
        collection_name=config.MILVUS_COLLECTION,
        data=[q["dense"]],
        limit=config.HYBRID_TOP_K,
        output_fields=["text", "source"],
        search_params={"metric_type": "COSINE"},
    )
    if not results or not results[0]:
        return []
    hits = results[0]

    # ③ 用 Reranker 精排（不可用就按召回顺序取 Top-K）
    candidates = [h["entity"]["text"] for h in hits]
    ranked = _rerank(query, candidates, top_k)
    if ranked is None:
        ranked = [(i, 1.0) for i in range(min(top_k, len(candidates)))]

    out = []
    for idx, score in ranked:
        h = hits[idx]
        out.append(
            {
                "text": h["entity"]["text"],
                "source": h["entity"]["source"],
                "score": round(score, 4),
            }
        )
    return out


def search_crawler_collection(query, top_k=5):
    """兜底检索：在成员A 爬虫流水线写入的 hacker_news 集合里做稠密检索。"""
    from app.rag.embedding import get_embedding_model
    from app.rag.milvus import get_milvus_manager

    embedding = get_embedding_model().encode([query])[0]
    manager = get_milvus_manager()
    return manager.search(embedding, limit=top_k)
