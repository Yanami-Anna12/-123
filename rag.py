"""RAG（检索增强生成）核心：把文档变成可检索的知识库。

完整流程（一条流水线）：
    文档文件 → 解析成文本 → 切分成片段 → 向量化(Dense+Sparse) → 存进 Milvus
    提问     → 向量化 → 混合检索(Dense+Sparse+RRF) → 重排(Reranker) → 返回 Top-K

外部包：
- 文档解析：pypdf / python-docx
- 文本切割：langchain 的 RecursiveCharacterTextSplitter
- 向量化：FlagEmbedding 的 BGEM3FlagModel（本地 BGE-M3，一次出 Dense+Sparse）
- 重排：sentence-transformers 的 CrossEncoder（本地 BGE-Reranker）
- 存储检索：pymilvus（Milvus）
"""
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pymilvus import AnnSearchRequest, DataType, MilvusClient, RRFRanker

import config

# 模型懒加载缓存：第一次真正用的时候才加载，之后复用（CPU 上加载一次要 20 多秒）
_embedding_model = None
_reranker_model = None


# ============ 1. 文档解析：文件 → 纯文本 ============

def _load_one(path):
    """读取单个文档，返回 langchain 的 Document（正文 + 来源信息）。"""
    p = Path(path)
    ext = p.suffix.lower()
    if ext in {".md", ".markdown", ".txt"}:
        # 纯文本，直接读出来
        text = p.read_text(encoding="utf-8", errors="ignore")
    elif ext == ".pdf":
        # PDF 用 pypdf 抽取每一页的文字
        from pypdf import PdfReader

        text = "\n\n".join(page.extract_text() or "" for page in PdfReader(str(p)).pages)
    elif ext == ".docx":
        # Word 用 python-docx 抽取段落文字
        import docx

        text = "\n".join(para.text for para in docx.Document(str(p)).paragraphs)
    else:
        raise ValueError(f"暂不支持的文件格式: {ext}")
    return Document(page_content=text, metadata={"source": p.name})


def load_documents(paths):
    """批量读取多个文档。"""
    return [_load_one(p) for p in paths]


# ============ 2. 文本切割：文本 → 片段（chunk） ============

def split_documents(documents):
    """用 langchain 的切割器切分文档（中文按句末标点切，段间留重叠）。"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        # 按这些分隔符从大到小尝试切：段落 → 换行 → 句末标点 → 逗号 → 空格 → 字符
        separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
    )
    return splitter.split_documents(documents)


# ============ 3. 向量化：文本 → Dense + Sparse 两种向量 ============

def get_embedding_model():
    """懒加载 BGE-M3 模型（CPU）。第一次调用才加载，之后复用。"""
    global _embedding_model
    if _embedding_model is None:
        from FlagEmbedding import BGEM3FlagModel

        _embedding_model = BGEM3FlagModel(
            config.EMBEDDING_MODEL_PATH,
            use_fp16=False,  # CPU 上用 32 位精度
            device="cpu",
        )
    return _embedding_model


def _embed_texts(texts):
    """把一批文本变成向量。每条返回 {"dense": [1024个float], "sparse": {词id: 权重}}。"""
    model = get_embedding_model()
    out = model.encode(
        texts,
        batch_size=8,
        return_dense=True,  # 要稠密向量（管“意思像不像”）
        return_sparse=True,  # 要稀疏向量（管“关键词对不对得上”）
        return_colbert_vecs=False,
    )
    dense = out["dense_vecs"]
    lexical = out["lexical_weights"]

    result = []
    for i in range(len(texts)):
        # 稀疏向量是 {词id: 权重}，把 key 从字符串转成整数（Milvus 要求整数）
        sparse = {int(k): float(v) for k, v in lexical[i].items()}
        result.append({"dense": list(dense[i]), "sparse": sparse})
    return result


# ============ 4. 重排：Reranker 对候选片段做二次精排 ============

def get_reranker_model():
    """懒加载 BGE-Reranker 模型（CPU）。"""
    global _reranker_model
    if _reranker_model is None:
        from sentence_transformers import CrossEncoder

        _reranker_model = CrossEncoder(config.RERANKER_MODEL_PATH, device="cpu")
    return _reranker_model


def _rerank(query, candidates, top_k):
    """对候选片段重新打分排序，返回 [(候选下标, 分数), ...]，分数从高到低。"""
    model = get_reranker_model()
    ranked = model.rank(
        query, candidates,
        top_k=min(top_k, len(candidates)),
        convert_to_numpy=True,
    )
    return [(int(item["corpus_id"]), float(item["score"])) for item in ranked]


# ============ 5. Milvus：建表 / 写入 / 混合检索 ============

def _client():
    """连接 Milvus。"""
    return MilvusClient(uri=config.MILVUS_URI)


def _init_collection():
    """建 Milvus 集合（如果不存在）。字段：id / text / source / dense_vector / sparse_vector。"""
    c = _client()
    if c.has_collection(config.MILVUS_COLLECTION):
        return c
    schema = c.create_schema(auto_id=True, enable_dynamic_field=False)
    schema.add_field("id", DataType.INT64, is_primary=True)
    schema.add_field("text", DataType.VARCHAR, max_length=65535)
    schema.add_field("source", DataType.VARCHAR, max_length=512)
    schema.add_field("dense_vector", DataType.FLOAT_VECTOR, dim=1024)  # 稠密向量
    schema.add_field("sparse_vector", DataType.SPARSE_FLOAT_VECTOR)  # 稀疏向量
    index = c.prepare_index_params()
    index.add_index(field_name="dense_vector", index_type="AUTOINDEX", metric_type="COSINE")
    index.add_index(field_name="sparse_vector", index_type="SPARSE_INVERTED_INDEX", metric_type="IP")
    c.create_collection(config.MILVUS_COLLECTION, schema=schema, index_params=index)
    c.load_collection(config.MILVUS_COLLECTION)
    return c


def ingest(paths):
    """文档入库：解析 → 切割 → 向量化 → 写入 Milvus，返回写入的片段数。"""
    chunks = split_documents(load_documents(paths))
    if not chunks:
        return 0
    embeds = _embed_texts([c.page_content for c in chunks])
    c = _init_collection()
    rows = [
        {
            "text": ch.page_content,
            "source": ch.metadata.get("source", ""),
            "dense_vector": e["dense"],
            "sparse_vector": e["sparse"],
        }
        for ch, e in zip(chunks, embeds)
    ]
    c.insert(config.MILVUS_COLLECTION, data=rows)
    c.flush(config.MILVUS_COLLECTION)
    return len(rows)


def search(query, top_k=None):
    """检索三步：① 向量化问题 ② 混合检索(Dense+Sparse+RRF) 召回 ③ 重排取 Top-K。"""
    top_k = top_k or config.TOP_K

    # ① 把问题变成 Dense + Sparse 两种向量
    q = _embed_texts([query])[0]

    # ② 混合检索：dense、sparse 各搜一路，再用 RRF 把两路排名融合成一个总排名
    c = _client()
    dense_req = AnnSearchRequest(
        data=[q["dense"]],
        anns_field="dense_vector",
        param={"metric_type": "COSINE"},
        limit=config.HYBRID_TOP_K,
    )
    sparse_req = AnnSearchRequest(
        data=[q["sparse"]],
        anns_field="sparse_vector",
        param={"metric_type": "IP"},
        limit=config.HYBRID_TOP_K,
    )
    results = c.hybrid_search(
        collection_name=config.MILVUS_COLLECTION,
        reqs=[dense_req, sparse_req],
        ranker=RRFRanker(k=60),  # RRF：融合两路排名的算法
        limit=config.HYBRID_TOP_K,
        output_fields=["text", "source"],
    )
    if not results:
        return []
    hits = results[0]  # 只有一个查询，取第一个查询的结果列表

    # ③ 用 Reranker 对候选片段做二次精排，取 Top-K
    candidates = [h["entity"]["text"] for h in hits]
    ranked = _rerank(query, candidates, top_k)

    out = []
    for idx, score in ranked:
        h = hits[idx]
        out.append(
            {
                "text": h["entity"]["text"],
                "source": h["entity"]["source"],
                "score": round(score, 4),
            }
        )
    return out
